#!/usr/bin/env python3
"""
Construye el libro completo en Word usando solo Python (sin pandoc).

Pasos:
  1. Renderiza todos los diagramas .mmd a .png con mmdc (mermaid-cli).
  2. Une los capítulos de capitulos/ en un único Markdown.
  3. Elimina los bloques ```mermaid ... ``` (las imágenes PNG ya están justo antes).
  4. Convierte el Markdown unificado a Word con python-docx.

Requisitos:
  - mmdc      (npm install -g @mermaid-js/mermaid-cli)
  - python-docx  (pip install python-docx)
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
CAPITULOS_DIR = PROJECT_ROOT / "capitulos"
IMAGES_DIR = PROJECT_ROOT / "assets" / "imagenes"
OUTPUT_DIR = PROJECT_ROOT / "output"
UNIFIED_MD = OUTPUT_DIR / "libro-completo.md"
OUTPUT_DOCX = OUTPUT_DIR / "libro-completo.docx"

MERMAID_BLOCK_RE = re.compile(r"```mermaid\s*\n.*?\n```\s*\n?", re.DOTALL)
DETAILS_BLOCK_RE = re.compile(r"<details>\s*.*?</details>\s*\n?", re.DOTALL | re.IGNORECASE)


# ---------------------------------------------------------------------------
# LaTeX -> texto Unicode (para fórmulas matemáticas)
# ---------------------------------------------------------------------------

GREEK = {
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\epsilon": "ε", r"\varepsilon": "ε", r"\zeta": "ζ", r"\eta": "η",
    r"\theta": "θ", r"\vartheta": "ϑ", r"\iota": "ι", r"\kappa": "κ",
    r"\lambda": "λ", r"\mu": "μ", r"\nu": "ν", r"\xi": "ξ",
    r"\pi": "π", r"\rho": "ρ", r"\sigma": "σ", r"\tau": "τ",
    r"\upsilon": "υ", r"\phi": "φ", r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
    r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
    r"\Xi": "Ξ", r"\Pi": "Π", r"\Sigma": "Σ", r"\Phi": "Φ",
    r"\Psi": "Ψ", r"\Omega": "Ω",
}

SYMBOLS = {
    r"\cdot": "·", r"\times": "×", r"\div": "÷", r"\pm": "±", r"\mp": "∓",
    r"\leq": "≤", r"\geq": "≥", r"\neq": "≠", r"\approx": "≈",
    r"\equiv": "≡", r"\sim": "∼", r"\infty": "∞", r"\partial": "∂",
    r"\nabla": "∇", r"\sum": "Σ", r"\prod": "∏", r"\int": "∫",
    r"\rightarrow": "→", r"\leftarrow": "←", r"\Rightarrow": "⇒",
    r"\Leftarrow": "⇐", r"\to": "→", r"\mapsto": "↦",
    r"\in": "∈", r"\notin": "∉", r"\subset": "⊂", r"\supset": "⊃",
    r"\cup": "∪", r"\cap": "∩", r"\forall": "∀", r"\exists": "∃",
    r"\langle": "⟨", r"\rangle": "⟩",
    r"\dots": "…", r"\ldots": "…", r"\cdots": "⋯",
    r"\ll": "≪", r"\gg": "≫", r"\propto": "∝",
}

CALLIGRAPHIC = {
    "A": "𝒜", "B": "ℬ", "C": "𝒞", "D": "𝒟", "E": "ℰ", "F": "ℱ",
    "G": "𝒢", "H": "ℋ", "I": "ℐ", "J": "𝒥", "K": "𝒦", "L": "ℒ",
    "M": "ℳ", "N": "𝒩", "O": "𝒪", "P": "𝒫", "Q": "𝒬", "R": "ℛ",
    "S": "𝒮", "T": "𝒯", "U": "𝒰", "V": "𝒱", "W": "𝒲", "X": "𝒳",
    "Y": "𝒴", "Z": "𝒵",
}

def _strip_braces(s: str) -> str:
    s = s.strip()
    while s.startswith("{") and s.endswith("}"):
        depth = 0
        ok = True
        for i, ch in enumerate(s):
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0 and i != len(s) - 1:
                    ok = False
                    break
        if not ok:
            break
        s = s[1:-1].strip()
    return s


def _read_brace_group(text: str, i: int) -> tuple[str, int]:
    """Lee {...} balanceado a partir de text[i]; devuelve (contenido, nuevo_i)."""
    if i >= len(text) or text[i] != "{":
        # Token simple: una letra o un caracter
        if i < len(text):
            return text[i], i + 1
        return "", i
    depth = 0
    start = i
    while i < len(text):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start + 1:i], i + 1
        i += 1
    return text[start + 1:], len(text)


def latex_preprocess(latex: str) -> str:
    """
    Aplica reemplazos LaTeX → Unicode SIN procesar aún los _ y ^.
    El resultado contiene letras griegas, símbolos y \\frac ya expandido,
    pero conserva los _ y ^ para que el renderizador los convierta en
    subíndices/superíndices reales de Word.
    """
    s = latex.strip()

    # \frac{a}{b} -> (a) / (b) (recursivo, balanceado)
    def replace_frac(text: str) -> str:
        out = []
        i = 0
        while i < len(text):
            if text.startswith(r"\frac", i):
                j = i + len(r"\frac")
                num, j = _read_brace_group(text, j)
                den, j = _read_brace_group(text, j)
                out.append(f"({replace_frac(num)}) / ({replace_frac(den)})")
                i = j
            else:
                out.append(text[i])
                i += 1
        return "".join(out)

    s = replace_frac(s)

    # \mathcal{X} -> letra caligráfica Unicode
    def replace_mathcal(m: re.Match) -> str:
        return "".join(CALLIGRAPHIC.get(c, c) for c in m.group(1))

    s = re.sub(r"\\mathcal\{([^}]*)\}", replace_mathcal, s)
    s = re.sub(r"\\mathbb\{([^}]*)\}", lambda m: m.group(1), s)
    s = re.sub(r"\\mathbf\{([^}]*)\}", lambda m: m.group(1), s)
    s = re.sub(r"\\text\{([^}]*)\}", lambda m: m.group(1), s)
    s = re.sub(r"\\operatorname\{([^}]*)\}", lambda m: m.group(1), s)

    # Funciones comunes (\log -> log, etc.)
    for fn in ("exp", "log", "ln", "sin", "cos", "tan", "max", "min", "arg", "clip"):
        s = s.replace(rf"\{fn}", fn)

    # Griegas y símbolos (orden: más largos primero para evitar solapes)
    for token in sorted(list(GREEK) + list(SYMBOLS), key=len, reverse=True):
        s = s.replace(token, GREEK.get(token, SYMBOLS.get(token, token)))

    # Limpieza
    s = s.replace(r"\,", " ").replace(r"\;", " ").replace(r"\!", "")
    s = s.replace(r"\left", "").replace(r"\right", "")
    s = s.replace(r"\mid", " | ")
    s = re.sub(r"\\\\", "\n", s)  # \\ -> salto de línea
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def latex_to_tokens(latex: str) -> list[tuple[str, str]]:
    """
    Convierte una fórmula LaTeX a una lista de tokens (texto, estilo)
    donde estilo ∈ {'normal', 'sub', 'super'}.

    Los subíndices y superíndices se emiten como tokens separados para
    que el renderizador pueda aplicar formato de run real de Word
    (font.subscript / font.superscript), no pseudo-Unicode.
    """
    s = latex_preprocess(latex)
    return _tokenize_scripts(s, base_style="normal")


def _tokenize_scripts(text: str, base_style: str) -> list[tuple[str, str]]:
    """Tokeniza `text` respetando _ y ^, propagando `base_style` al contenido normal."""
    tokens: list[tuple[str, str]] = []
    buf = ""
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]
        if ch in ("_", "^"):
            if buf:
                tokens.append((buf, base_style))
                buf = ""
            content, j = _read_brace_group(text, i + 1)
            content = _strip_braces(content)
            inner_style = "sub" if ch == "_" else "super"
            # Recursión: el contenido puede tener más _/^ anidados
            for t, st in _tokenize_scripts(content, base_style=inner_style):
                # Si el padre es sub/super y el hijo también, mantenemos el del hijo
                tokens.append((t, st))
            i = j
        else:
            buf += ch
            i += 1
    if buf:
        tokens.append((buf, base_style))
    return tokens


def render_math_runs(paragraph, latex: str, *, base_size: Pt | None = None) -> None:
    """Añade los runs de una fórmula a `paragraph` con sub/superscript reales."""
    for text, style in latex_to_tokens(latex):
        if not text:
            continue
        r = paragraph.add_run(text)
        r.italic = True
        r.font.name = "Cambria Math"
        if base_size is not None:
            r.font.size = base_size
        if style == "sub":
            r.font.subscript = True
        elif style == "super":
            r.font.superscript = True


# ---------------------------------------------------------------------------
# Paso 1: render Mermaid -> PNG
# ---------------------------------------------------------------------------

def render_mermaid_diagrams() -> int:
    mmdc = shutil.which("mmdc")
    if not mmdc:
        sys.exit(
            "ERROR: 'mmdc' no está instalado.\n"
            "  -> npm install -g @mermaid-js/mermaid-cli"
        )

    mmd_files = sorted(IMAGES_DIR.glob("*.mmd"))
    if not mmd_files:
        print(f"No se encontraron archivos .mmd en {IMAGES_DIR}")
        return 0

    rendered = 0
    for mmd in mmd_files:
        png = mmd.with_suffix(".png")
        if png.exists() and png.stat().st_mtime >= mmd.stat().st_mtime:
            print(f"  [skip] {mmd.name}")
            continue
        print(f"  [render] {mmd.name} -> {png.name}")
        result = subprocess.run(
            [mmdc, "-i", str(mmd), "-o", str(png), "-b", "transparent", "-w", "1200"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            sys.exit(f"ERROR renderizando {mmd.name}:\n{result.stderr}")
        rendered += 1

    print(f"Diagramas renderizados: {rendered}")
    return rendered


# ---------------------------------------------------------------------------
# Paso 2: unificar capítulos en un único Markdown
# ---------------------------------------------------------------------------

def build_unified_markdown() -> Path:
    md_files = sorted(CAPITULOS_DIR.glob("*.md"))
    if not md_files:
        sys.exit(f"No se encontraron capítulos en {CAPITULOS_DIR}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    parts: list[str] = []

    for md in md_files:
        content = md.read_text(encoding="utf-8")
        content = DETAILS_BLOCK_RE.sub("", content)
        content = MERMAID_BLOCK_RE.sub("", content)
        parts.append("")
        parts.append(content)
        parts.append("")

    UNIFIED_MD.write_text("\n".join(parts), encoding="utf-8")
    print(f"Markdown unificado: {UNIFIED_MD}")
    return UNIFIED_MD


# ---------------------------------------------------------------------------
# Paso 3: Markdown -> DOCX (parser propio, suficiente para nuestros capítulos)
# ---------------------------------------------------------------------------

INLINE_PATTERN = re.compile(
    r"(\*\*\*(.+?)\*\*\*)"            # 1,2  bold-italic
    r"|(\*\*(.+?)\*\*)"               # 3,4  bold
    r"|(\*(.+?)\*)"                   # 5,6  italic
    r"|(__(.+?)__)"                   # 7,8  bold (underscore)
    r"|(`([^`]+)`)"                   # 9,10 inline code
    r"|(\[([^\]]+)\]\(([^)]+)\))"     # 11,12 text 13 url
    r"|(\$([^\$\n]+)\$)"              # 14,15 inline math
)


def add_runs(paragraph, text: str) -> None:
    """Añade runs con formato inline (negrita, cursiva, código, links)."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group(1):  # bold-italic
            r = paragraph.add_run(m.group(2))
            r.bold = True
            r.italic = True
        elif m.group(3) or m.group(7):  # bold
            r = paragraph.add_run(m.group(4) or m.group(8))
            r.bold = True
        elif m.group(5):  # italic
            r = paragraph.add_run(m.group(6))
            r.italic = True
        elif m.group(9):  # inline code
            r = paragraph.add_run(m.group(10))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif m.group(11):  # link [text](url)
            r = paragraph.add_run(m.group(12))
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            r.underline = True
        elif m.group(14):  # inline math $...$
            render_math_runs(paragraph, m.group(15), base_size=Pt(11))

        pos = m.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[Imagen no encontrada: {image_path.name}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(6))
    except Exception as exc:
        run.text = f"[Error insertando {image_path.name}: {exc}]"
        return
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)


_BOOKMARK_ID = [1000]


def slugify(text: str) -> str:
    """Convierte un texto a un id válido de bookmark de Word (sin acentos, sin espacios)."""
    import unicodedata
    norm = unicodedata.normalize("NFKD", text)
    norm = "".join(c for c in norm if not unicodedata.combining(c))
    norm = re.sub(r"[^\w]+", "_", norm).strip("_")
    return ("h_" + norm)[:40] or f"h_{_BOOKMARK_ID[0]}"


def add_heading_with_bookmark(doc: Document, text: str, level: int, bookmark_id: str) -> None:
    """Añade un heading con un bookmark interno para poder linkar a él."""
    h = doc.add_heading(level=min(level, 4))
    # Bookmark start
    bm_id = _BOOKMARK_ID[0]
    _BOOKMARK_ID[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), bookmark_id)
    h._p.append(start)
    add_runs(h, text)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    h._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    """Añade un hipervínculo interno (a un bookmark) dentro de un párrafo."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def collect_headings(lines: list[str]) -> list[tuple[int, str, str]]:
    """Recorre las líneas, ignora bloques de código/math, y devuelve [(nivel, texto, slug)]."""
    headings: list[tuple[int, str, str]] = []
    seen_slugs: dict[str, int] = {}
    in_code = False
    in_math = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if stripped.startswith("$$"):
            # toggle si abre y cierra en distintas líneas
            if stripped.endswith("$$") and len(stripped) > 4:
                continue
            in_math = not in_math
            continue
        if in_math:
            continue
        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            # Quitar formato inline (negrita, código...) para el slug y display
            display = re.sub(r"[`*_]", "", title)
            base = slugify(display)
            count = seen_slugs.get(base, 0)
            seen_slugs[base] = count + 1
            slug = base if count == 0 else f"{base}_{count}"
            headings.append((level, title, slug))
    return headings


def add_toc(doc: Document, headings: list[tuple[int, str, str]]) -> None:
    """Inserta un índice manual con hyperlinks internos clicables."""
    h = doc.add_heading("Tabla de Contenidos", level=1)

    for level, title, slug in headings:
        if level > 4:
            continue
        display = re.sub(r"[`*_]", "", title)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
        p.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(p, display, slug)

    # Salto de página después del índice
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_math_block(doc: Document, latex: str) -> None:
    """Inserta una fórmula display centrada con sub/superscripts reales."""
    # El preprocesado puede generar saltos de línea (\\\\ en LaTeX)
    for line in latex_preprocess(latex).split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Re-tokenizamos la línea ya preprocesada
        for text, style in _tokenize_scripts(line, base_style="normal"):
            if not text:
                continue
            r = p.add_run(text)
            r.italic = True
            r.font.name = "Cambria Math"
            r.font.size = Pt(12)
            if style == "sub":
                r.font.subscript = True
            elif style == "super":
                r.font.superscript = True


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            add_runs(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True


IMAGE_LINE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
UL_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
OL_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_cover_page(doc: Document) -> None:
    """Crea una portada profesional para el libro."""
    # Espaciado superior
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FINE-TUNING DE LLMs")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r.font.name = "Calibri Light"

    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 40)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("De la Teoría a la Práctica")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    r.font.name = "Calibri Light"
    r.italic = True

    # Descripción
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(
        "Guía completa de fine-tuning de modelos de lenguaje:\n"
        "Transformers, SFT, LoRA, QLoRA, RLHF y GRPO"
    )
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.name = "Calibri"

    # Espaciado inferior
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Fuente
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Basado en las Finetuning Sessions de The Neural Maze")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.italic = True

    # Salto de página
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada profesional
    add_cover_page(doc)

    # Recolectar headings y construir TOC clicable
    headings = collect_headings(lines)
    add_toc(doc, headings)
    heading_iter = iter(headings)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Línea en blanco
        if not stripped:
            i += 1
            continue

        # Regla horizontal
        if stripped in ("---", "***", "___"):
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # Bloque math display $$...$$
        if stripped.startswith("$$"):
            # Caso una sola línea: $$...$$
            if stripped.endswith("$$") and len(stripped) > 4:
                add_math_block(doc, stripped[2:-2])
                i += 1
                continue
            # Multilínea
            buf = [stripped[2:]]
            i += 1
            while i < n and not lines[i].strip().endswith("$$"):
                buf.append(lines[i])
                i += 1
            if i < n:
                last = lines[i].strip()
                buf.append(last[:-2])
                i += 1
            add_math_block(doc, "\n".join(buf))
            continue

        # Bloque de código
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            i += 1
            buf: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # cierre
            add_code_block(doc, "\n".join(buf), language)
            continue

        # Imagen en su propia línea
        m_img = IMAGE_LINE_RE.match(stripped)
        if m_img:
            alt, url = m_img.group(1), m_img.group(2)
            img_path = (PROJECT_ROOT / url).resolve()
            add_image(doc, img_path, alt)
            i += 1
            continue

        # Encabezado
        m_h = HEADING_RE.match(stripped)
        if m_h:
            level = len(m_h.group(1))
            title = m_h.group(2).strip()
            try:
                _, _, slug = next(heading_iter)
            except StopIteration:
                slug = slugify(title)
            add_heading_with_bookmark(doc, title, level, slug)
            i += 1
            continue

        # Tabla: línea actual con '|' y siguiente con guiones
        if "|" in stripped and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]):
            header = split_table_row(stripped)
            i += 2  # saltar separador
            rows = [header]
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Blockquote
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            add_runs(p, " ".join(buf))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Listas (no anidadas; las anidadas se aplanan)
        m_ul = UL_RE.match(line)
        m_ol = OL_RE.match(line)
        if m_ul or m_ol:
            ordered = m_ol is not None
            style_name = "List Number" if ordered else "List Bullet"
            while i < n:
                m_ul = UL_RE.match(lines[i])
                m_ol = OL_RE.match(lines[i])
                if not (m_ul or m_ol):
                    break
                content = (m_ol or m_ul).group(2)
                p = doc.add_paragraph(style=style_name)
                add_runs(p, content)
                i += 1
            continue

        # Párrafo normal: acumular líneas hasta blanco/elemento especial
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_strip = nxt.strip()
            if not nxt_strip:
                break
            if (nxt_strip.startswith("#")
                or nxt_strip.startswith("```")
                or nxt_strip.startswith("$$")
                or nxt_strip.startswith(">")
                or IMAGE_LINE_RE.match(nxt_strip)
                or UL_RE.match(nxt) or OL_RE.match(nxt)
                or nxt_strip in ("---", "***", "___")
                or ("|" in nxt_strip and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]))):
                break
            buf.append(nxt_strip)
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(buf))

    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size / 1024
    print(f"Documento generado: {docx_path} ({size_kb:.1f} KB)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("== Paso 1: Renderizando diagramas Mermaid ==")
    render_mermaid_diagrams()

    print("\n== Paso 2: Unificando capítulos ==")
    unified = build_unified_markdown()

    print("\n== Paso 3: Generando Word ==")
    markdown_to_docx(unified, OUTPUT_DOCX)

    print("\nListo.")


if __name__ == "__main__":
    main()
